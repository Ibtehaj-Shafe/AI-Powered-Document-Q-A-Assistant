import os, io, uuid
from dotenv import load_dotenv
from datetime import datetime, timezone
from sqlalchemy.orm import Session
from Backend.models.document import Document
from Backend.models import UserStats
from PyPDF2 import PdfReader
import docx
from sentence_transformers import SentenceTransformer
from pinecone import Pinecone, ServerlessSpec

# Load environment variables
load_dotenv()
key = os.getenv("db_key")          # Pinecone API key
env = os.getenv("reg")             # Pinecone region
index_name = os.getenv("index_name")

# Initialize Pinecone client
pc = Pinecone(api_key=key)
index = pc.Index(index_name)

# Load Sentence Transformer model
embedder = SentenceTransformer("all-MiniLM-L6-v2")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


def process_upload(file_content: str, filename: str, user_id: int, db: Session) -> Document:
    """
    Handles the full upload workflow:
    1. Save document metadata in DB
    2. Update user's stats
    3. Chunk text, embed locally, and upsert into Pinecone
    """

    # 1. Save document metadata
    doc = Document(
        filename=filename,
        user_id=user_id,
        upload_date=datetime.now(timezone.utc)
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # 2. Update stats
    stats = db.query(UserStats).filter(UserStats.user_id == user_id).first()
    if not stats:
        stats = UserStats(user_id=user_id, files_uploaded_count=1)
        db.add(stats)
    else:
        stats.files_uploaded_count += 1
    db.commit()

    # 3. Chunk text + embed + upsert to Pinecone
    chunks = chunk_text(file_content)
    vectors = []
    for chunk in chunks:
        embedding = embedder.encode(chunk).tolist()
        vectors.append((
            str(uuid.uuid4()),
            embedding,
            {
                "user_id": user_id,
                "document_id": doc.id,
                "filename": filename,
                "chunk": chunk
            }
        ))

    index.upsert(vectors)

    return doc